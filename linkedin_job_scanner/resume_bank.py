from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from .models import ResumeDocument
from .text_utils import DEFAULT_SKILL_PHRASES, cosine_similarity, extract_keywords, normalize_text, phrase_in_text


SECTION_HEADINGS = {
    "professional summary",
    "summary",
    "core competencies",
    "core competencies and skills",
    "skills",
    "technical skills",
    "professional experience",
    "work experience",
    "experience",
    "internships and research",
    "strategic projects",
    "strategic analytics and research projects",
    "education",
    "certifications",
    "projects",
}


@dataclass
class ResumeBank:
    root: Path
    documents: list[ResumeDocument]
    profile_text: str
    profile_keywords: list[str]

    @classmethod
    def build(cls, root: str | Path, config: dict[str, Any]) -> "ResumeBank":
        root_path = Path(root).expanduser().resolve()
        cache_path = Path(config["output_dir"]) / "data" / "resume_profile_cache.json"
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        docx_paths = sorted(_candidate_resume_paths(root_path, config))
        trusted_root = Path(str(config.get("trusted_resume_root", ""))).expanduser()
        if trusted_root.exists():
            docx_paths.extend(sorted(_candidate_resume_paths(trusted_root.resolve(), config)))
        signature = {str(path): path.stat().st_mtime for path in docx_paths if path.exists()}

        cached = _load_cache(cache_path)
        if cached and cached.get("signature") == signature:
            docs = [ResumeDocument.from_dict(item) for item in cached.get("documents", [])]
            profile_text = str(cached.get("profile_text", ""))
            keywords = list(cached.get("profile_keywords", []))
            return cls(root=root_path, documents=docs, profile_text=profile_text, profile_keywords=keywords)

        docs: list[ResumeDocument] = []
        for path in docx_paths:
            try:
                paragraphs = read_docx_paragraphs(path)
            except Exception as exc:
                print(f"Skipping unreadable resume {path}: {exc}")
                continue
            text = "\n".join(paragraphs)
            if len(text.strip()) < 300:
                continue
            docs.append(ResumeDocument(path=str(path), text=text, paragraphs=paragraphs))

        profile_text = "\n".join(_dedupe_paragraphs(p for doc in docs for p in doc.paragraphs))
        keywords = extract_keywords(profile_text, DEFAULT_SKILL_PHRASES, top_n=120)
        cache_path.write_text(
            json.dumps(
                {
                    "signature": signature,
                    "documents": [doc.to_dict() for doc in docs],
                    "profile_text": profile_text,
                    "profile_keywords": keywords,
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        return cls(root=root_path, documents=docs, profile_text=profile_text, profile_keywords=keywords)

    def best_resume_for_job(self, job_text: str, preferred_paths: list[str] | None = None) -> ResumeDocument:
        if not self.documents:
            raise RuntimeError("No resume DOCX files were found. Check resume_root in config.json.")

        preferred_bonus: dict[str, float] = {}
        for order, rel_path in enumerate(preferred_paths or []):
            path = (self.root / rel_path).resolve()
            if path.exists():
                preferred_bonus[str(path)] = max(0.02, 0.25 - order * 0.04)

        best_doc = self.documents[0]
        best_score = -1.0
        for doc in self.documents:
            score = cosine_similarity(job_text, doc.text)
            score += preferred_bonus.get(doc.path, 0.0)
            if score > best_score:
                best_score = score
                best_doc = doc
        return best_doc

    def evidence_for_keyword(self, keyword: str, limit: int = 3) -> list[str]:
        hits: list[str] = []
        for paragraph in _dedupe_paragraphs(p for doc in self.documents for p in doc.paragraphs):
            if len(paragraph) < 30:
                continue
            if phrase_in_text(keyword, paragraph):
                hits.append(paragraph)
            if len(hits) >= limit:
                break
        return hits

    def supported_keywords(self, keywords: list[str]) -> list[str]:
        return [kw for kw in keywords if phrase_in_text(kw, self.profile_text)]


def read_docx_paragraphs(path: str | Path) -> list[str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Run: pip install -r job_scanner/requirements.txt") from exc

    doc = Document(str(path))
    paragraphs: list[str] = []
    for paragraph in doc.paragraphs:
        text = _clean_para(paragraph.text)
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean_para(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                paragraphs.append(text)
    return paragraphs


def is_section_heading(text: str) -> bool:
    cleaned = normalize_text(text)
    if cleaned in SECTION_HEADINGS:
        return True
    return text.isupper() and len(text.split()) <= 7 and len(text) > 3


def _candidate_resume_paths(root: Path, config: dict[str, Any]) -> list[Path]:
    exclude_terms = [normalize_text(term) for term in config.get("exclude_resume_name_terms", [])]
    excluded_parts = {".linkedin_profile", "outputs", "__pycache__"}
    candidates: list[Path] = []
    for path in root.rglob("*.docx"):
        rel_parts = set(path.relative_to(root).parts)
        if rel_parts & excluded_parts:
            continue
        if "job_scanner" in rel_parts:
            continue
        name = normalize_text(path.name)
        rel = normalize_text(str(path.relative_to(root)))
        if any(term and term in rel for term in exclude_terms):
            continue
        if path.name.startswith("~$"):
            continue
        if "resume" in name or "cv" in name or "nikhil" in name:
            candidates.append(path)
    return candidates


def _clean_para(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).strip()


def _dedupe_paragraphs(paragraphs: Any) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for paragraph in paragraphs:
        cleaned = _clean_para(str(paragraph))
        key = normalize_text(cleaned)
        if not cleaned or key in seen:
            continue
        seen.add(key)
        result.append(cleaned)
    return result


def _load_cache(path: Path) -> dict[str, Any] | None:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None
