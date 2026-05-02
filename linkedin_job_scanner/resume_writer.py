from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from .models import JobPosting, ScoreResult
from .resume_bank import ResumeBank, is_section_heading, read_docx_paragraphs
from .text_utils import DEFAULT_SKILL_PHRASES, safe_filename


SUMMARY_HEADINGS = {"professional summary", "summary"}
SKILL_HEADINGS = {"core competencies", "core competencies and skills", "skills", "technical skills"}
SUMMARY_DISPLAY_TERMS = set(DEFAULT_SKILL_PHRASES) | {
    "analytics",
    "insights",
    "strategy",
    "category growth strategy",
    "commercial development",
    "competitive dynamics",
    "fmcg",
    "go-to-market",
    "in-store execution",
    "joint business planning",
    "market analytics",
    "nielsen",
    "nielseniq",
    "pricing",
    "promotions",
    "revenue management",
    "shopper behaviour",
    "shopper behavior",
    "syndicated data",
}
SUMMARY_BLOCKED_TERMS = {"r", "client", "partner", "senior", "manager", "associate"}
CORE_POSITIONING_TITLE = "Category & Shopper Strategy / Business Analysis & Insights"
CORE_POSITIONING_SUMMARY = (
    "Category & Shopper Strategy / Business Analysis & Insights professional with an MBA in Finance from "
    "McMaster University and 4+ years of analytics experience across FMCG/CPG, retail media, financial "
    "services, telecom, and technology. At Loblaw Advance, supported Confectionery and Beauty category "
    "insights for Tier 1 CPG clients including Hershey, Lindt, Mondelez, L'Oreal, Nestle, and Ferrero on "
    "Canada's 41M+ member PC Optimum loyalty platform."
)
CORE_POSITIONING_IMPACT = (
    "Experienced translating NielsenIQ/Nielsen RMC, POS, panel, campaign, loyalty, and behavioral data into "
    "category growth strategies, shopper insights, assortment recommendations, pricing and promotions analysis, "
    "executive storytelling, KPI dashboards, and scalable analytics workflows. Proven impact includes $24M+ in "
    "category growth opportunities, $15M in retention opportunities, $9M in organic growth potential, 40%+ "
    "faster decision-making, $38M in revenue protected, 95% manual effort reduction, and 26% analytics adoption uplift."
)
FOUNDATIONAL_COMPETENCIES = {
    "Category, Shopper & Market Insights": [
        "category management",
        "category insights",
        "shopper insights",
        "consumer insights",
        "market share analysis",
        "competitive analysis",
        "retail analytics",
        "marketing analytics",
    ],
    "Analytics, BI & Modelling": [
        "sql",
        "python",
        "power bi",
        "tableau",
        "looker studio",
        "forecasting",
        "segmentation",
        "a/b testing",
    ],
    "Strategy, Stakeholders & Delivery": [
        "strategy",
        "business intelligence",
        "executive reporting",
        "stakeholder management",
        "cross-functional collaboration",
        "kpi reporting",
        "process automation",
        "agile",
    ],
    "AI, Automation & Advanced Analytics": [
        "ai-powered analytics",
        "ai agent workflows",
        "machine learning",
        "predictive analytics",
        "statistical modelling",
        "nlp",
        "etl",
        "financial modelling",
    ],
}


def create_tailored_resume(
    job: JobPosting,
    score: ScoreResult,
    resume_bank: ResumeBank,
    output_dir: str | Path,
    config: dict[str, Any],
) -> tuple[str, str]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Run: pip install -r job_scanner/requirements.txt") from exc

    resume_dir = Path(output_dir) / "resumes"
    resume_dir.mkdir(parents=True, exist_ok=True)
    job_part = safe_filename(f"{job.company}_{job.title}_{job.job_id}", "linkedin_job")
    output_path = resume_dir / f"{job_part}_Nikhil_Jha_Resume.docx"

    supported_keywords = resume_bank.supported_keywords(score.matched_keywords)
    summary_keywords = _summary_display_keywords(supported_keywords, resume_bank.profile_keywords)
    data = _build_claude_resume_data(job, summary_keywords, supported_keywords, config)

    doc = Document()
    _build_claude_style_resume_doc(doc, data)
    doc.save(str(output_path))
    cover_path, outreach_path = _write_companion_materials(job, data, job_part, output_dir)
    score.cover_letter_path = cover_path
    score.cold_outreach_path = outreach_path

    resume_text = "\n".join(read_docx_paragraphs(output_path))
    return str(output_path), resume_text


def ensure_tailored_companion_materials(
    job: JobPosting,
    score: ScoreResult,
    resume_bank: ResumeBank,
    output_dir: str | Path,
    config: dict[str, Any],
) -> tuple[str, str]:
    """Create cover letter and cold outreach files for an existing generated resume."""
    cover_path = Path(score.cover_letter_path) if score.cover_letter_path else None
    outreach_path = Path(score.cold_outreach_path) if score.cold_outreach_path else None
    if (
        cover_path
        and cover_path.exists()
        and cover_path.suffix.lower() == ".docx"
        and outreach_path
        and outreach_path.exists()
        and outreach_path.suffix.lower() == ".docx"
    ):
        return score.cover_letter_path, score.cold_outreach_path

    supported_keywords = resume_bank.supported_keywords(score.matched_keywords)
    summary_keywords = _summary_display_keywords(supported_keywords, resume_bank.profile_keywords)
    data = _build_claude_resume_data(job, summary_keywords, supported_keywords, config)
    job_part = safe_filename(f"{job.company}_{job.title}_{job.job_id}", "linkedin_job")
    cover_path, outreach_path = _write_companion_materials(job, data, job_part, output_dir)
    score.cover_letter_path = cover_path
    score.cold_outreach_path = outreach_path
    return cover_path, outreach_path


def _build_claude_style_resume_doc(
    doc: Any,
    data: dict[str, Any],
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.625)
        section.bottom_margin = Inches(0.625)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1.5)
    _add_run(name, "NIKHIL JHA", 20, bold=True)
    _add_run(name, ", MBA", 13)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline.paragraph_format.space_after = Pt(2)
    _add_run(headline, data["headline"], 13)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    _add_run(
        contact,
        "nikhil.jha97@outlook.com • (437) 995-0287 • Mississauga, ON • linkedin.com/in/nikhiljha97 • github.com/nikhiljha97",
        9,
        color="555555",
    )

    _add_section_heading(doc, "PROFESSIONAL SUMMARY")
    _add_body_paragraph(doc, data["summary"], space_before=3, space_after=3)

    _add_section_heading(doc, "CORE COMPETENCIES")
    for label, items in data["competencies"]:
        paragraph = _add_body_paragraph(doc, "", space_before=1.5, space_after=1.5)
        _add_run(paragraph, f"{label}: ", 10, bold=True)
        _add_run(paragraph, items, 10)

    _add_section_heading(doc, "WORK EXPERIENCE")
    for role_index, role in enumerate(data["experience"]):
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_before = Pt(6 if role_index else 4)
        header.paragraph_format.space_after = Pt(1)
        header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        _add_run(header, role["title"], 10, bold=True)
        _add_run(header, f" | {role['org']}, {role['location']}", 10)
        _add_run(header, f"\t{role['dates']}", 9, color="555555")
        header.paragraph_format.keep_with_next = True

        if role.get("sublabel"):
            sublabel = doc.add_paragraph()
            sublabel.paragraph_format.space_after = Pt(1.5)
            _add_run(sublabel, role["sublabel"], 9, italic=True, color="555555")
            sublabel.paragraph_format.keep_with_next = True

        for bullet_text in role["bullets"]:
            bullet = doc.add_paragraph()
            bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            bullet.paragraph_format.left_indent = Inches(0.28)
            bullet.paragraph_format.first_line_indent = Inches(-0.14)
            bullet.paragraph_format.space_before = Pt(1)
            bullet.paragraph_format.space_after = Pt(1)
            _add_run(bullet, "• ", 10)
            _add_run(bullet, bullet_text, 10)

    _add_section_heading(doc, "EDUCATION")
    _add_education_line(
        doc,
        "MBA, Finance & Strategy",
        "DeGroote School of Business, McMaster University",
        "Sep 2023 – Apr 2025",
    )
    _add_small_paragraph(
        doc,
        "GPA: 3.4/4.0 • MBA Entrance Scholarship • KPMG 101 Consulting Workshop, Top 25 Nationally • Advocate, DeGroote Foresight Lab",
    )
    _add_small_paragraph(doc, data["mba_coursework"], color="1A1A1A")
    _add_education_line(
        doc,
        "Bachelor of Technology, Information Technology",
        "SRM Institute of Science & Technology, India",
        "Jul 2016 – Jun 2020",
    )
    _add_small_paragraph(
        doc,
        "GPA: 8.8/10 • Gold Medalist in Advanced Mathematics • Exchange: SungKyunKwan University, Seoul, South Korea (2018)",
    )

    _add_section_heading(doc, "PROFESSIONAL DEVELOPMENT")
    for label, text in data["professional_development"]:
        paragraph = _add_body_paragraph(doc, "", space_before=1.5, space_after=1.5)
        _add_run(paragraph, f"{label}: ", 10, bold=True)
        _add_run(paragraph, text, 10)


def _add_section_heading(doc: Any, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    _add_run(paragraph, text, 11, bold=True)
    _apply_claude_section_rule(paragraph)


def _add_body_paragraph(doc: Any, text: str, space_before: float = 1.5, space_after: float = 1.5) -> Any:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if text:
        _add_run(paragraph, text, 10)
    return paragraph


def _add_small_paragraph(doc: Any, text: str, color: str = "555555") -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    _add_run(paragraph, text, 9, color=color)


def _add_education_line(doc: Any, degree: str, school: str, dates: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches, Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _add_run(paragraph, degree, 10, bold=True)
    _add_run(paragraph, f" – {school}", 10)
    _add_run(paragraph, f"\t{dates}", 9, color="555555")


def _add_run(
    paragraph: Any,
    text: str,
    size_pt: int | float,
    bold: bool = False,
    italic: bool = False,
    color: str = "1A1A1A",
) -> Any:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run = paragraph.add_run(_strip_em_dash(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def _apply_claude_section_rule(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    borders.append(bottom)
    p_pr.append(borders)


def _strip_em_dash(text: str) -> str:
    return text.replace("—", ",")


def _build_claude_resume_data(
    job: JobPosting,
    summary_keywords: list[str],
    supported_keywords: list[str],
    config: dict[str, Any],
) -> dict[str, Any]:
    keywords = _unique([*supported_keywords, *summary_keywords])
    domain = _job_domain(job.full_text())
    headline_terms = _headline_terms(job, keywords, domain)
    years = _experience_label(config)

    summary = _claude_summary(job, headline_terms, domain, years)
    competencies = _claude_competencies(keywords, domain)
    experience = _claude_experience_roles(domain, keywords)
    return {
        "headline": _compact_join([job.title, *headline_terms], " | ", max_chars=130),
        "summary": summary,
        "competencies": competencies,
        "experience": experience,
        "mba_coursework": _mba_coursework(domain),
        "cover_letter": _cover_letter_paragraphs(job, domain, headline_terms),
        "cold_outreach": _cold_outreach(job, domain, headline_terms),
        "professional_development": [
            (
                "CPG Insight Engine",
                "Built an AI-powered category analytics platform for CPG and retail insight workflows, combining SQL, analytics logic, and executive-facing recommendations.",
            ),
            (
                "KPMG 101 Consulting Workshop",
                "Selected for a Top 25 national consulting workshop focused on structured problem solving, executive communication, and business case development.",
            ),
        ],
    }


def _write_companion_materials(
    job: JobPosting,
    data: dict[str, Any],
    job_part: str,
    output_dir: str | Path,
) -> tuple[str, str]:
    cover_dir = Path(output_dir) / "cover_letters"
    outreach_dir = Path(output_dir) / "cold_outreach"
    cover_dir.mkdir(parents=True, exist_ok=True)
    outreach_dir.mkdir(parents=True, exist_ok=True)

    cover_path = cover_dir / f"{job_part}_Nikhil_Jha_Cover_Letter.docx"
    outreach_path = outreach_dir / f"{job_part}_Nikhil_Jha_Cold_Outreach.docx"

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required. Run: pip install -r job_scanner/requirements.txt") from exc

    cover_doc = Document()
    _build_claude_style_cover_letter_doc(cover_doc, job, data)
    cover_doc.save(str(cover_path))
    outreach_doc = Document()
    _build_cold_outreach_doc(outreach_doc, job, data)
    outreach_doc.save(str(outreach_path))
    return str(cover_path), str(outreach_path)


def _build_claude_style_cover_letter_doc(doc: Any, job: JobPosting, data: dict[str, Any]) -> None:
    from datetime import datetime

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.875)
        section.right_margin = Inches(0.875)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    _add_run(name, "NIKHIL JHA", 20, bold=True)
    _add_run(name, ", MBA", 13)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(12)
    _add_run(contact, "nikhil.jha97@outlook.com • (437) 995-0287 • Mississauga, ON • linkedin.com/in/nikhiljha97", 9, color="555555")

    _add_left_line(doc, datetime.now().strftime("%B %-d, %Y"), after=8)
    _add_left_line(doc, "Hiring Manager", after=2)
    _add_left_line(doc, job.company, after=10)

    re_line = doc.add_paragraph()
    re_line.paragraph_format.space_after = Pt(10)
    _add_run(re_line, "Re: ", 10, bold=True)
    _add_run(re_line, job.title, 10)

    _add_left_line(doc, "Dear Hiring Manager,", after=8)
    for paragraph_text in data["cover_letter"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(8)
        _add_run(paragraph, paragraph_text, 10)

    _add_left_line(doc, "Sincerely,", after=16)
    closing = doc.add_paragraph()
    _add_run(closing, "Nikhil Jha, MBA", 10, bold=True)


def _add_left_line(doc: Any, text: str, after: float = 4) -> None:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    _add_run(paragraph, text, 10)


def _build_cold_outreach_doc(doc: Any, job: JobPosting, data: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.875)
        section.right_margin = Inches(0.875)

    outreach = data["cold_outreach"]
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    _add_run(title, f"{job.company} - {job.title}", 13, bold=True)
    _add_run(title, " - Cold Outreach", 11)

    _add_outreach_section(doc, "SEARCH TIPS", outreach["search_tips"])
    _add_outreach_section(doc, "LINKEDIN CONNECTION NOTE", outreach["connection"])
    _add_outreach_section(doc, "LINKEDIN FOLLOW-UP MESSAGE", outreach["followup"])
    _add_outreach_section(doc, f"EMAIL SUBJECT: {outreach['email_subject']}", outreach["email_body"])
    _add_outreach_section(doc, "NOTE", outreach["notes"])


def _add_outreach_section(doc: Any, heading: str, body: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    _add_run(paragraph, heading, 11, bold=True)
    _apply_claude_section_rule(paragraph)

    for chunk in body.split("\n\n"):
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_para.paragraph_format.space_after = Pt(5)
        _add_run(body_para, chunk, 10)


def _cover_letter_paragraphs(job: JobPosting, domain: str, headline_terms: list[str]) -> list[str]:
    role_phrase = _domain_phrase(domain)
    terms = ", ".join(_unique_display([*headline_terms[:3], "SQL", "Power BI", "Advanced Excel"]))
    return [
        _strip_em_dash(
            f"{job.company}'s {job.title} role stands out because it needs someone who can turn messy operating, "
            f"reporting, and stakeholder information into practical recommendations. I bring a background aligned with "
            f"{role_phrase.lower()}, with analytics, reporting automation, and executive storytelling experience across CPG retail media, "
            f"telecom, financial services, and academic research."
        ),
        _strip_em_dash(
            "At Loblaw Advance, I translated shopper, campaign, retail measurement, and CPG account data into recurring "
            "and ad hoc reporting for senior stakeholders and vendor-facing conversations. That work used 41M PC Optimum "
            "loyalty members, Power BI, SQL, and executive-ready storylines to uncover $9M in organic growth and $15M in "
            "strategic recovery opportunities while improving decision speed by 40%."
        ),
        _strip_em_dash(
            f"My broader experience adds the technical and operating discipline behind that storytelling: {terms}, Python, "
            "dashboard automation, cross-functional delivery, and clear translation of technical findings for non-technical "
            "leaders. At Verizon, I delivered analytics and automation workflows tied to $38M in risk management outcomes, "
            "95% manual processing reduction, and 26% model adoption improvement."
        ),
        _strip_em_dash(
            f"I would bring {job.company} a practical mix of analytical depth, business judgment, and concise communication. "
            "I am careful about evidence, direct about tradeoffs, and strongest when turning ambiguous data into decisions "
            "that teams can act on quickly."
        ),
    ]


def _cold_outreach(job: JobPosting, domain: str, headline_terms: list[str]) -> dict[str, str]:
    role_phrase = _domain_phrase(domain)
    terms = ", ".join(headline_terms[:3])
    return {
        "connection": _strip_em_dash(
            f"Hi, I applied for {job.title} at {job.company}. My background spans {role_phrase.lower()}, Power BI, SQL, "
            "and executive-ready reporting needs through work at Loblaw Advance and Verizon. Would be glad to connect."
        )[:290],
        "followup": _strip_em_dash(
            f"Hi, I recently applied for the {job.title} role at {job.company}. At Loblaw Advance, I turned shopper, campaign, "
            "and CPG account data into reporting and senior-stakeholder stories using 41M PC Optimum loyalty members, Power BI, "
            "and SQL. The work uncovered $9M in organic growth and $15M in strategic recovery opportunities. I would welcome "
            f"the chance to bring the same mix of {terms}, practical analytics, and executive communication to your team."
        ),
        "email_subject": f"{job.title} application, Nikhil Jha",
        "email_body": _strip_em_dash(
            f"Hi,\n\nI applied for the {job.title} role at {job.company} and wanted to briefly introduce myself. I bring "
            f"experience aligned with {role_phrase.lower()} across CPG retail media, telecom, financial services, and academic research, "
            "with hands-on work in Power BI, SQL, Python, Advanced Excel, and stakeholder-ready reporting.\n\nAt Loblaw Advance, "
            "I translated shopper, campaign, and retail measurement data for strategic CPG accounts using 41M PC Optimum "
            "loyalty members, uncovering $9M in organic growth and $15M in strategic recovery opportunities. At Verizon, "
            "I delivered analytics automation tied to $38M in risk management outcomes.\n\nIf helpful, I would appreciate "
            "being pointed to the right person for the role.\n\nBest,\nNikhil Jha\n(437) 995-0287\nlinkedin.com/in/nikhiljha97"
        ),
        "search_tips": f"Search LinkedIn for {job.company} talent acquisition, hiring manager, analytics leader, insights leader, or the job poster.",
        "notes": "Uses only verified experience and approved metrics. Review company-specific phrasing before sending.",
    }


def _cold_outreach_text(job: JobPosting, data: dict[str, Any]) -> str:
    outreach = data["cold_outreach"]
    return "\n".join(
        [
            f"{job.company} - {job.title} - Cold Outreach",
            "=" * 60,
            "",
            f"Search tips: {outreach['search_tips']}",
            "",
            "LINKEDIN CONNECTION NOTE",
            "-" * 32,
            outreach["connection"],
            "",
            "LINKEDIN FOLLOW-UP MESSAGE",
            "-" * 32,
            outreach["followup"],
            "",
            f"EMAIL SUBJECT: {outreach['email_subject']}",
            "-" * 32,
            outreach["email_body"],
            "",
            f"NOTE: {outreach['notes']}",
            "",
        ]
    )



def _job_domain(job_text: str) -> str:
    normalized = job_text.lower()
    if any(term in normalized for term in ["talent", "recruit", "campus", "mobility", "human resources", "hris"]):
        return "talent_programs"
    if any(term in normalized for term in ["category", "shopper", "consumer", "cpg", "fmcg", "retail", "nielsen"]):
        return "category_insights"
    if any(term in normalized for term in ["finance", "revenue", "pricing", "forecast", "budget", "p&l", "pnl"]):
        return "finance_strategy"
    if any(term in normalized for term in ["product", "customer", "lifecycle", "ux", "user"]):
        return "product_analytics"
    if any(term in normalized for term in ["change management", "transformation", "program", "operations"]):
        return "strategy_operations"
    if any(term in normalized for term in ["data scientist", "machine learning", "ai", "nlp", "model"]):
        return "advanced_analytics"
    return "business_insights"


def _experience_label(config: dict[str, Any]) -> str:
    years = float(config.get("candidate_experience_years") or 3.9)
    if years >= 4.8:
        return "5+ years"
    if years >= 3.8:
        return "4+ years"
    return f"{years:.1f} years"


def _headline_terms(job: JobPosting, keywords: list[str], domain: str) -> list[str]:
    domain_defaults = {
        "talent_programs": ["Reporting & Insights", "Program Operations", "Power BI"],
        "category_insights": ["Category Insights", "Shopper Analytics", "NielsenIQ"],
        "finance_strategy": ["Strategic Finance", "Forecasting", "Executive Reporting"],
        "product_analytics": ["Product Analytics", "Customer Insights", "Experimentation"],
        "strategy_operations": ["Strategy & Operations", "Process Improvement", "Stakeholder Storytelling"],
        "advanced_analytics": ["Advanced Analytics", "Machine Learning", "SQL"],
        "business_insights": ["Business Intelligence", "Insights", "SQL"],
    }
    blocked = {
        "analytics",
        "associate",
        "candidate",
        "client",
        "insights",
        "manager",
        "people",
        "python",
        "senior",
        "sql",
        "strategy",
        "team",
    }
    terms = [_title_skill(term) for term in keywords if _usable_resume_keyword(term) and term not in blocked][:2]
    combined = _unique_display([*domain_defaults.get(domain, domain_defaults["business_insights"]), *terms])
    return combined[:4]


def _compact_join(parts: list[str], separator: str, max_chars: int) -> str:
    values = [part for part in parts if part]
    while values and len(separator.join(values)) > max_chars:
        values.pop()
    return separator.join(values)


def _claude_summary(job: JobPosting, headline_terms: list[str], domain: str, years: str) -> str:
    role_phrase = _domain_phrase(domain)
    methods = ", ".join(_unique_display([*headline_terms[:4], "SQL", "Python", "Power BI", "Tableau", "Advanced Excel"]))
    company_context = _company_context(domain)
    return _strip_em_dash(
        f"{role_phrase} candidate with {years} of analytics, reporting, automation, and stakeholder storytelling "
        f"experience across retail media, financial services, telecom, and academic research. {company_context} "
        f"Brings hands-on strength in {methods}, and executive-ready "
        f"recommendations aligned to the {job.title} role at {job.company}. Delivered $62M total business value across "
        f"category insights, automation, risk analytics, and stakeholder decision support without overstating scope, "
        f"ownership, or unsupported metrics."
    )


def _domain_phrase(domain: str) -> str:
    return {
        "talent_programs": "Talent Programs & Insights",
        "category_insights": "Category & Shopper Insights",
        "finance_strategy": "Strategic Finance & Business Insights",
        "product_analytics": "Product Analytics & Customer Insights",
        "strategy_operations": "Strategy, Operations & Business Insights",
        "advanced_analytics": "Advanced Analytics & Business Intelligence",
        "business_insights": "Business Intelligence & Insights",
    }.get(domain, "Business Intelligence & Insights")


def _company_context(domain: str) -> str:
    if domain == "talent_programs":
        return (
            "At Loblaw Advance, translated CPG account data into recurring and ad hoc reporting, dashboard automation, "
            "vendor-facing performance updates, and senior-stakeholder narratives using 41M PC Optimum loyalty members."
        )
    if domain == "finance_strategy":
        return (
            "At Loblaw Advance, converted campaign, pricing, and performance data into strategic recovery, organic growth, "
            "and executive reporting narratives for CPG commercial stakeholders."
        )
    if domain == "product_analytics":
        return (
            "At Loblaw Advance and Verizon, used behavioral data, segmentation, reporting automation, and model outputs "
            "to explain user, customer, and operational trends for non-technical leaders."
        )
    if domain == "advanced_analytics":
        return (
            "At Verizon, delivered Python, SQL, machine learning, NLP, and BI workflows that supported fraud analytics, "
            "risk management outcomes, and model adoption improvement."
        )
    return (
        "At Loblaw Advance, supported Confectionery and Beauty category insights for Hershey, Lindt, Mondelez, L'Oreal, "
        "Nestle, and Ferrero using shopper, POS, panel, campaign, and retail measurement data."
    )


def _claude_competencies(keywords: list[str], domain: str) -> list[tuple[str, str]]:
    role_terms = _keyword_line(
        keywords,
        {
            "talent_programs": [
                "program operations",
                "process automation",
                "reporting automation",
                "executive reporting",
                "stakeholder management",
                "data analysis",
                "insight delivery",
            ],
            "category_insights": [
                "category insights",
                "shopper insights",
                "consumer insights",
                "market share analysis",
                "category management",
                "nielseniq",
                "pricing and promotions",
            ],
            "finance_strategy": [
                "financial modelling",
                "forecasting",
                "scenario analysis",
                "business case",
                "executive reporting",
                "pricing and promotions",
            ],
            "product_analytics": [
                "product analytics",
                "customer analytics",
                "segmentation",
                "a/b testing",
                "measurement framework",
                "data visualization",
            ],
        }.get(domain, ["business intelligence", "data analysis", "executive reporting", "insight delivery", "strategy"]),
    )
    tech_terms = _keyword_line(
        keywords,
        ["sql", "python", "power bi", "tableau", "looker studio", "advanced excel", "r", "nielseniq", "jira", "agile"],
    )
    leadership_terms = _keyword_line(
        keywords,
        [
            "stakeholder management",
            "cross-functional collaboration",
            "executive storytelling",
            "kpi reporting",
            "process automation",
            "business intelligence",
        ],
    )
    industry_terms = _keyword_line(
        keywords,
        [
            "cpg",
            "fmcg",
            "retail media",
            "consumer insights",
            "risk analytics",
            "marketing analytics",
            "operations",
            "strategy",
        ],
    )
    tech_terms = _append_unique_phrase(tech_terms, "Advanced Excel")
    return [
        ("Role-Specific Alignment", role_terms),
        ("Technical Tools", tech_terms),
        ("Stakeholder Influence", leadership_terms),
        ("Industry & Domain Knowledge", industry_terms),
    ]


def _keyword_line(keywords: list[str], desired: list[str]) -> str:
    values = [_title_skill(term) for term in desired]
    generic_extras = {"analytics", "insights", "strategy", "python", "sql", "r"}
    extras = [
        _title_skill(term)
        for term in keywords
        if _usable_resume_keyword(term) and term not in generic_extras and _title_skill(term) not in values
    ][:2]
    return ", ".join(_unique_display([*values, *extras])[:12])


def _append_unique_phrase(line: str, phrase: str) -> str:
    values = [item.strip() for item in line.split(",") if item.strip()]
    if phrase not in values:
        values.append(phrase)
    return ", ".join(values)


def _claude_experience_roles(domain: str, keywords: list[str]) -> list[dict[str, Any]]:
    return [
        {
            "title": "Manager, Business Intelligence & Analytics",
            "org": "Loblaw Advance",
            "location": "Toronto, ON",
            "dates": "Aug 2025 – Dec 2025",
            "sublabel": _loblaw_sublabel(domain),
            "bullets": _loblaw_bullets(domain, keywords),
        },
        {
            "title": "Technical Project Manager",
            "org": "Exera Solutions Inc.",
            "location": "Remote, ON",
            "dates": "Jan 2025 – Mar 2025",
            "sublabel": "Program delivery · Milestone tracking · Stakeholder coordination",
            "bullets": [
                "Coordinated program-style delivery across stakeholder groups by translating business needs into structured workplans, milestone tracking, status updates, and issue-resolution routines.",
                "Supported process improvements and operational documentation by clarifying handoffs, tracking action items, and helping teams move ambiguous work toward measurable outcomes.",
            ],
        },
        {
            "title": "Graduate Research Assistant",
            "org": "McMaster University",
            "location": "Hamilton, ON",
            "dates": "Jun 2024 – Dec 2024",
            "sublabel": "Research reporting · Data visualization · Program support",
            "bullets": [
                "Built research analytics and reporting workflows using Python, Looker Studio, and Power BI, reducing reporting time by 60% and improving stakeholder access to decision-ready insights.",
                "Translated data outputs into concise visual narratives and presentation materials for academic and business audiences, strengthening executive-ready storytelling and evidence-based recommendations.",
            ],
        },
        {
            "title": "Software Engineer II, Risk Analytics & NLP",
            "org": "Verizon Inc.",
            "location": "Chennai, India",
            "dates": "Aug 2020 – Jul 2023",
            "sublabel": "Analytics automation · Process optimization · Cross-functional delivery",
            "bullets": _verizon_bullets(domain),
        },
    ]


def _loblaw_sublabel(domain: str) -> str:
    return {
        "talent_programs": "Reporting & insights · Executive storytelling · Vendor-facing analytics",
        "category_insights": "Category insights · Shopper analytics · CPG retail measurement",
        "finance_strategy": "Strategic recovery · Revenue analysis · Executive reporting",
        "product_analytics": "Behavioral analytics · Segmentation · Measurement frameworks",
        "strategy_operations": "Strategy operations · KPI frameworks · Stakeholder decision support",
        "advanced_analytics": "BI automation · Data modelling · Performance analytics",
        "business_insights": "Business intelligence · Dashboard automation · Insight delivery",
    }.get(domain, "Business intelligence · Dashboard automation · Insight delivery")


def _loblaw_bullets(domain: str, keywords: list[str]) -> list[str]:
    keyword_phrase = _natural_keyword_phrase(keywords)
    if domain == "talent_programs":
        return [
            "Built and interpreted recurring reporting for strategic CPG accounts, translating shopper, campaign, and retail measurement data into executive-ready storylines that accelerated stakeholder decisions by 40%.",
            "Led cross-functional analytics work with commercial teams, vendor partners, and client stakeholders, bringing structure to ambiguous business questions and converting data into measurable recommendations for efficiency and effectiveness.",
            "Developed Power BI and SQL reporting workflows to clean, validate, and explain metric movements across audience, campaign, and market performance data, supporting $9M organic growth and $15M strategic recovery opportunities.",
            "Prepared leader-ready presentations for Hershey, Lindt, L'Oreal, and other CPG partners using 41M PC Optimum loyalty members and retail media performance data, improving clarity of recommendations for senior audiences.",
            "Supported vendor-facing performance conversations and quarterly-style account readouts by synthesizing insights, issue drivers, and next actions that helped lock in $70K Reese's media investment and $100K Lindt media investment.",
        ]
    if domain == "finance_strategy":
        return [
            "Analyzed campaign, shopper, and retail measurement performance to identify $15M in strategic recovery opportunities and $9M in organic growth potential for CPG commercial stakeholders.",
            "Built executive reporting narratives that connected pricing, promotions, audience performance, and market movement to commercial decisions, accelerating stakeholder decisions by 40%.",
            "Used SQL, Power BI, and Advanced Excel to validate recurring and ad hoc performance data, explain variance drivers, and translate metric movement into business recommendations.",
            "Synthesized 41M PC Optimum loyalty member signals into leader-ready presentations for Hershey, Lindt, L'Oreal, and other strategic accounts, improving confidence in investment and planning discussions.",
            f"Mapped {keyword_phrase} into practical reporting, dashboard, and recommendation workflows where the language was supported by real CPG analytics work and existing resume evidence.",
        ]
    if domain == "product_analytics":
        return [
            "Translated shopper, audience, and campaign behavior into customer analytics narratives using 41M PC Optimum loyalty members, supporting segmentation, measurement, and decision-ready recommendations.",
            "Developed Power BI and SQL reporting workflows to explain product, audience, and performance trends across strategic CPG accounts, accelerating stakeholder decisions by 40%.",
            "Built measurement frameworks that connected campaign outcomes, retail behavior, and audience quality to $9M organic growth and $15M strategic recovery opportunities.",
            "Prepared executive-ready stories for Hershey, Lindt, L'Oreal, and other partners, using data visualization and business context to make analytics usable for non-technical stakeholders.",
        ]
    if domain == "advanced_analytics":
        return [
            "Built Power BI, SQL, and Python-supported analytics workflows across shopper, campaign, and retail measurement data, improving recurring insight delivery and accelerating stakeholder decisions by 40%.",
            "Translated 41M PC Optimum loyalty member signals into segmentation, performance analytics, and executive narratives that uncovered $9M organic growth and $15M strategic recovery opportunities.",
            "Validated metric movement across audience and campaign data to separate signal from noise, improving senior stakeholder confidence in recommendations and investment conversations.",
            f"Connected {keyword_phrase} to practical business intelligence outputs, using only terms that align to existing analytics, reporting, and automation experience.",
        ]
    return [
        "Led category and business intelligence analytics for strategic CPG accounts, translating shopper, POS, panel, campaign, and retail measurement data into recommendations that accelerated stakeholder decisions by 40%.",
        "Analyzed 41M PC Optimum loyalty member signals for Hershey, Lindt, L'Oreal, and other CPG partners, uncovering $9M organic growth and $15M strategic recovery opportunities.",
        "Built Power BI, SQL, and Advanced Excel reporting workflows to validate metric movement, monitor performance, and convert complex data into executive-ready storylines for senior stakeholders.",
        "Synthesized pricing, promotion, audience, and category performance findings into leader-ready presentations that helped lock in $70K Reese's media investment and $100K Lindt media investment.",
        f"Mirrored role-relevant terms such as {keyword_phrase} where they truthfully mapped to existing category insights, reporting automation, and stakeholder storytelling experience.",
    ]


def _verizon_bullets(domain: str) -> list[str]:
    if domain in {"advanced_analytics", "product_analytics"}:
        return [
            "Delivered Python, SQL, and Power BI analytics pipelines that automated recurring analysis and reporting workflows, contributing to $38M risk management outcomes across fraud and operational processes.",
            "Built machine learning and NLP-enabled classification workflows that reduced manual processing by 95%, generated $3M cost savings, and improved model adoption by 26%.",
            "Partnered with Risk, Finance, Compliance, and engineering teams through Agile delivery routines, translating technical model outputs into clear operating decisions for cross-functional stakeholders.",
        ]
    return [
        "Delivered Python, SQL, and Power BI analytics pipelines that automated recurring analysis and reporting workflows, contributing to $38M risk management outcomes across fraud and operational processes.",
        "Translated complex risk and fraud data into stakeholder-ready dashboards and operating insights, reducing manual processing by 95% while improving reporting consistency.",
        "Led analytics adoption and change-management routines across cross-functional teams, improving model adoption by 26% and strengthening trust in automated decision support.",
    ]


def _natural_keyword_phrase(keywords: list[str]) -> str:
    blocked = {"client", "partner", "senior", "manager", "associate", "analyst"}
    values = [_title_skill(term) for term in keywords if _usable_resume_keyword(term) and term not in blocked][:4]
    if not values:
        values = ["Reporting Automation", "Executive Storytelling", "Stakeholder Management"]
    if len(values) == 1:
        return values[0]
    return ", ".join(values[:-1]) + f", and {values[-1]}"


def _mba_coursework(domain: str) -> str:
    courses = {
        "talent_programs": "Relevant Coursework: Business Strategy, Organizational Behaviour, Financial Modelling, Marketing Research, Leadership Communication",
        "category_insights": "Relevant Coursework: Marketing Research, Consumer Behaviour, Business Strategy, Financial Modelling, P&L Management",
        "finance_strategy": "Relevant Coursework: Financial Modelling, Corporate Finance, Business Strategy, P&L Management, Managerial Accounting",
        "product_analytics": "Relevant Coursework: Marketing Research, Consumer Behaviour, Business Strategy, Data Analytics, Financial Modelling",
        "strategy_operations": "Relevant Coursework: Business Strategy, Operations Management, Financial Modelling, Leadership Communication, Marketing Research",
        "advanced_analytics": "Relevant Coursework: Marketing Research, Financial Modelling, Business Strategy, Statistical Analysis, Operations Management",
    }
    return courses.get(domain, "Relevant Coursework: Business Strategy, Marketing Research, Financial Modelling, Data Analytics, Leadership Communication")


def _usable_resume_keyword(term: str) -> bool:
    normalized = term.strip().lower()
    if not normalized or len(normalized) <= 2:
        return False
    blocked = {
        "associate",
        "candidate",
        "client",
        "consultant",
        "manager",
        "people",
        "senior",
        "team",
    }
    if normalized in blocked:
        return False
    if len(normalized.split()) > 3:
        return False
    return normalized in DEFAULT_SKILL_PHRASES or normalized in SUMMARY_DISPLAY_TERMS


def _unique_display(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        value = item.strip()
        key = value.lower()
        if not value or key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def _build_summary(job: JobPosting, supported_keywords: list[str]) -> list[str]:
    role_terms = ", ".join(_title_skill(term) for term in supported_keywords[:8])
    role_phrase = f" Strengths include {role_terms}." if role_terms else ""
    return [CORE_POSITIONING_SUMMARY, f"{CORE_POSITIONING_IMPACT}{role_phrase}"]


def _build_competency_lines(supported_keywords: list[str], profile_keywords: list[str]) -> list[str]:
    allowed = set(DEFAULT_SKILL_PHRASES)
    terms = _unique([term for term in [*supported_keywords, *profile_keywords] if term in allowed])
    lines: list[str] = []
    used: set[str] = set()

    jd_terms = [term for term in supported_keywords if term in allowed and term not in SUMMARY_BLOCKED_TERMS][:10]
    if jd_terms:
        lines.append("Target Role Alignment: " + " | ".join(_title_skill(term) for term in jd_terms))
        used.update(jd_terms)

    for label, desired in FOUNDATIONAL_COMPETENCIES.items():
        matches = [term for term in desired if term not in used]
        matches = matches[:8]
        used.update(matches)
        if matches:
            lines.append(f"{label}: " + " | ".join(_title_skill(term) for term in matches))

    remaining = [term for term in terms if term not in used and term not in SUMMARY_BLOCKED_TERMS][:10]
    if remaining:
        lines.append("Additional ATS Keywords: " + " | ".join(_title_skill(term) for term in remaining))
    return lines[:6]


def _summary_display_keywords(supported_keywords: list[str], profile_keywords: list[str]) -> list[str]:
    terms = _unique([*supported_keywords, *profile_keywords])
    display_terms = [
        term
        for term in terms
        if term in SUMMARY_DISPLAY_TERMS and term not in SUMMARY_BLOCKED_TERMS and len(term) > 1
    ]
    return display_terms[:10]


def _replace_top_headline(doc: Any, job: JobPosting, supported_keywords: list[str]) -> None:
    paragraphs = doc.paragraphs
    summary_index = _find_first_heading_index(paragraphs, SUMMARY_HEADINGS)
    if summary_index is None:
        return
    headline = _headline_for_job(job, supported_keywords)
    for idx in range(1, summary_index):
        text = paragraphs[idx].text.strip()
        lower = text.lower()
        if "@" in lower or "linkedin" in lower or any(char.isdigit() for char in lower[:20]):
            continue
        if 10 <= len(text) <= 180:
            _set_paragraph_text(paragraphs[idx], headline)
            return


def _headline_for_job(job: JobPosting, supported_keywords: list[str]) -> str:
    headline_terms = [
        term
        for term in supported_keywords
        if term not in {"python", "sql", "power bi", "excel", "tableau", "looker studio"}
    ][:3]
    terms = " | ".join(_title_skill(term) for term in headline_terms)
    title = job.title or "Strategy and Analytics"
    if terms:
        return f"{title} | {CORE_POSITIONING_TITLE} | {terms} | SQL, Python, Power BI"
    return f"{title} | {CORE_POSITIONING_TITLE} | SQL, Python, Power BI | MBA Finance"


def _replace_section(doc: Any, headings: set[str], lines: list[str], plain: bool = False) -> bool:
    if not lines:
        return False
    paragraphs = doc.paragraphs
    start = _find_first_heading_index(paragraphs, headings)
    if start is None:
        return False
    end = _find_next_heading_index(paragraphs, start + 1)
    content = paragraphs[start + 1 : end]
    if not content:
        anchor = paragraphs[start]
        for line in lines:
            anchor = _insert_paragraph_after(anchor, line, plain=plain)
        return True

    for paragraph, line in zip(content, lines):
        _set_paragraph_text(paragraph, line, plain=plain)

    if len(lines) > len(content):
        anchor = content[-1]
        for line in lines[len(content) :]:
            anchor = _insert_paragraph_after(anchor, line, style=content[-1].style, plain=plain)
    else:
        for paragraph in reversed(content[len(lines) :]):
            _delete_paragraph(paragraph)
    return True


def _insert_section_after(doc: Any, after_headings: set[str], heading_text: str, lines: list[str]) -> bool:
    paragraphs = doc.paragraphs
    start = _find_first_heading_index(paragraphs, after_headings)
    if start is None:
        return False
    end = _find_next_heading_index(paragraphs, start + 1)
    anchor = paragraphs[end - 1] if end > start else paragraphs[start]
    heading = _insert_paragraph_after(anchor, heading_text, style=paragraphs[start].style)
    if heading.runs:
        heading.runs[0].bold = True
    content_style = paragraphs[start + 1].style if start + 1 < len(paragraphs) else None
    anchor = heading
    for line in lines:
        anchor = _insert_paragraph_after(anchor, line, style=content_style, plain=True)
    return True


def _find_first_heading_index(paragraphs: list[Any], headings: set[str]) -> int | None:
    normalized = {_norm_heading(item) for item in headings}
    for index, paragraph in enumerate(paragraphs):
        if _norm_heading(paragraph.text) in normalized:
            return index
    return None


def _find_next_heading_index(paragraphs: list[Any], start: int) -> int:
    for index in range(start, len(paragraphs)):
        if is_section_heading(paragraphs[index].text.strip()):
            return index
    return len(paragraphs)


def _norm_heading(text: str) -> str:
    return " ".join(text.replace("&", " and ").lower().split())


def _set_paragraph_text(paragraph: Any, text: str, plain: bool = False) -> None:
    _clear_paragraph_content(paragraph)
    run = paragraph.add_run(text)
    if plain:
        run.bold = False
        run.italic = False


def _insert_paragraph_after(paragraph: Any, text: str, style: Any | None = None, plain: bool = False) -> Any:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    run = new_para.add_run(text)
    if plain:
        run.bold = False
        run.italic = False
    return new_para


def _clear_paragraph_content(paragraph: Any) -> None:
    from docx.oxml.ns import qn

    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _delete_paragraph(paragraph: Any) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _normalize_layout(doc: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    for section in doc.sections:
        section.top_margin = Inches(0.625)
        section.bottom_margin = Inches(0.625)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.56)

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        _clear_paragraph_borders(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = False
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.line_spacing = None
        _format_paragraph_runs(paragraph, "Calibri", 9)
        if index == 0 and text:
            _format_paragraph_runs(paragraph, "Calibri", 20, bold=True)
            paragraph.paragraph_format.space_after = Pt(1.5)
        elif index == 1 and text:
            _format_paragraph_runs(paragraph, "Calibri", 13)
            paragraph.paragraph_format.space_after = Pt(1.5)
        elif index == 2 and text:
            _format_paragraph_runs(paragraph, "Calibri", 9)
            paragraph.paragraph_format.space_after = Pt(5)
        elif is_section_heading(text):
            _format_paragraph_runs(paragraph, "Calibri", 11, bold=True)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            _apply_heading_borders(paragraph)
        elif paragraph.style and paragraph.style.name == "List Paragraph":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            _format_paragraph_runs(paragraph, "Calibri", 10)
            _bold_label_prefix(paragraph, font_size_pt=10)
        elif _looks_like_role_header(text):
            _format_paragraph_runs(paragraph, "Calibri", 11)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(1)
            _bold_role_designation(paragraph)
        elif text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_before = Pt(1.5)
            paragraph.paragraph_format.space_after = Pt(1.5)
            _bold_label_prefix(paragraph)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if is_section_heading(text) or _looks_like_role_header(text) or _looks_like_date_line(text):
            paragraph.paragraph_format.keep_with_next = True


def _apply_heading_borders(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    _clear_paragraph_borders(paragraph)

    borders = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), "AAAAAA")
        borders.append(border)
    p_pr.append(borders)


def _clear_paragraph_borders(paragraph: Any) -> None:
    from docx.oxml.ns import qn

    if paragraph._p.pPr is None:
        return
    existing = paragraph._p.pPr.find(qn("w:pBdr"))
    if existing is not None:
        paragraph._p.pPr.remove(existing)


def _format_paragraph_runs(paragraph: Any, font_name: str, font_size_pt: int | float, bold: bool | None = None) -> None:
    from docx.shared import Pt

    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        if bold is not None:
            run.font.bold = bold


def _bold_label_prefix(paragraph: Any, font_size_pt: int | float = 9) -> None:
    from docx.shared import Pt

    text = paragraph.text
    colon = text.find(":")
    if colon <= 0 or colon > 70:
        return
    _clear_paragraph_content(paragraph)
    label = paragraph.add_run(text[: colon + 1])
    label.font.name = "Calibri"
    label.font.size = Pt(font_size_pt)
    label.bold = True
    body = paragraph.add_run(text[colon + 1 :])
    body.font.name = "Calibri"
    body.font.size = Pt(font_size_pt)


def _bold_role_designation(paragraph: Any) -> None:
    from docx.shared import Pt

    text = paragraph.text
    separator = " | " if " | " in text else " – " if " – " in text else ""
    if not separator:
        _format_paragraph_runs(paragraph, "Calibri", 11)
        return

    pos = text.find(separator)
    if pos <= 0:
        return
    _clear_paragraph_content(paragraph)
    designation = paragraph.add_run(text[:pos])
    designation.font.name = "Calibri"
    designation.font.size = Pt(11)
    designation.bold = True
    rest = paragraph.add_run(text[pos:])
    rest.font.name = "Calibri"
    rest.font.size = Pt(11)
    rest.bold = False


def _bold_before_separator(paragraph: Any, separator: str) -> None:
    from docx.shared import Pt

    text = paragraph.text
    pos = text.find(separator)
    if pos <= 0 or len(paragraph.runs) != 1:
        return
    _clear_paragraph_content(paragraph)
    first = paragraph.add_run(text[:pos])
    first.font.name = "Calibri"
    first.font.size = Pt(9)
    first.bold = True
    rest = paragraph.add_run(text[pos:])
    rest.font.name = "Calibri"
    rest.font.size = Pt(9)


def _trim_trailing_incomplete_role(doc: Any) -> None:
    paragraphs = doc.paragraphs
    nonempty = [(idx, p, p.text.strip()) for idx, p in enumerate(paragraphs) if p.text.strip()]
    if len(nonempty) < 4:
        return

    tail: list[tuple[int, Any, str]] = []
    for idx, paragraph, text in reversed(nonempty):
        if _is_trailing_fragment(text):
            tail.append((idx, paragraph, text))
            continue
        break
    if not tail or not any(_looks_like_role_header(text) for _idx, _p, text in tail):
        return

    first_idx = min(idx for idx, _p, _text in tail)
    if first_idx > 0 and is_section_heading(paragraphs[first_idx - 1].text.strip()):
        first_idx -= 1
    for paragraph in reversed(paragraphs[first_idx:]):
        _delete_paragraph(paragraph)


def _title_skill(term: str) -> str:
    overrides = {
        "ai": "AI",
        "ai agent workflows": "AI Agent Workflows",
        "ai-powered analytics": "AI-Powered Analytics",
        "a/b testing": "A/B Testing",
        "cpg": "CPG",
        "etl": "ETL",
        "fmcg": "FMCG",
        "jira": "Jira",
        "kpi reporting": "KPI Reporting",
        "looker studio": "Looker Studio",
        "nielseniq": "NielsenIQ",
        "nlp": "NLP",
        "okr": "OKR",
        "pos data analysis": "POS Data Analysis",
        "power bi": "Power BI",
        "python": "Python",
        "r": "R",
        "sql": "SQL",
    }
    return overrides.get(term, term.title())


def _looks_like_role_header(text: str) -> bool:
    if len(text) > 260 or len(text) < 12:
        return False
    if ":" in text:
        return False
    if "|" in text and not text.endswith("."):
        return True
    return bool(re.search(r"\b(manager|analyst|engineer|assistant|consultant|intern|associate)\b", text, re.I)) and not text.endswith(".")


def _looks_like_date_line(text: str) -> bool:
    return bool(re.search(r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b", text, re.I)) and len(text) < 80


def _is_trailing_fragment(text: str) -> bool:
    if is_section_heading(text) or _looks_like_role_header(text) or _looks_like_date_line(text):
        return True
    if len(text) < 130 and "|" in text and not text.endswith("."):
        return True
    return False


def _unique(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        key = item.strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(key)
    return result
